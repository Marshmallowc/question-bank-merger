#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
题库合并工具
支持自定义Excel格式，将多个题库文件合并成一个统一的Word或Excel文档
"""
import pandas as pd
import os
import glob
import json
import argparse
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，无法生成Word文档")


class QuestionBankMerger:
    def __init__(self, config_path: str = "config/config.json"):
        """初始化题库合并工具"""
        self.config = self.load_config(config_path)
        self.merged_data = None

    def load_config(self, config_path: str) -> Dict:
        """加载配置文件"""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"配置文件 {config_path} 不存在，使用默认配置")
            return self.get_default_config()
        except json.JSONDecodeError as e:
            print(f"配置文件格式错误: {e}")
            return self.get_default_config()

    def get_default_config(self) -> Dict:
        """获取默认配置"""
        return {
            "excel_settings": {
                "has_header_row": True,
                "header_row_index": 1,  # 第2行（0-based为1）
                "data_start_row": 2,    # 第3行开始是数据
                "skip_description_row": True,
                "description_row_index": 0
            },
            "column_mapping": {
                "question_type": "题型",
                "question_text": "题干",
                "correct_answer": "正确答案",
                "analysis": "解析",
                "score": "分值",
                "difficulty": "难度系数",
                "options": ["选项A", "选项B", "选项C", "选项D", "选项E"]
            },
            "output_settings": {
                "excel_filename": "merged_questions.xlsx",
                "word_filename": "merged_questions.docx",
                "include_analysis": True,
                "include_difficulty": True
            },
            "file_patterns": [
                "*_习题导出.xlsx",
                "*questions*.xlsx",
                "*题库*.xlsx",
                "*.xlsx"
            ]
        }

    def detect_format(self, filepath: str) -> Dict:
        """自动检测Excel文件格式"""
        df = pd.read_excel(filepath, engine='openpyxl', header=None)

        # 检测结构
        format_info = {
            "has_header": True,
            "header_row": 1,
            "data_start": 2,
            "columns": []
        }

        # 查找包含"题型"的行作为表头
        for idx, row in df.iterrows():
            if idx >= 5:  # 只检查前5行
                break
            for cell in row:
                if pd.notna(cell) and '题型' in str(cell):
                    format_info["header_row"] = idx
                    format_info["data_start"] = idx + 1
                    return format_info

        return format_info

    def read_excel_file(self, filepath: str) -> pd.DataFrame:
        """读取Excel文件"""
        print(f"正在读取: {filepath}")

        try:
            # 根据配置读取文件
            if self.config["excel_settings"]["skip_description_row"]:
                # 跳过描述行的情况 - 先读取所有行再处理
                df_temp = pd.read_excel(filepath, engine='openpyxl', header=None)
                column_names = df_temp.iloc[self.config["excel_settings"]["header_row_index"]]
                data = df_temp[self.config["excel_settings"]["data_start_row"]:].copy()
                data.columns = column_names
                df = data
            else:
                # 直接读取
                df = pd.read_excel(
                    filepath,
                    engine='openpyxl',
                    header=self.config["excel_settings"]["header_row_index"]
                )

            # 重置索引
            df.reset_index(drop=True, inplace=True)

            # 清理数据
            df = self.clean_data(df)

            # 添加文件来源信息
            filename = Path(filepath).stem
            df["来源文件"] = filename

            print(f"  [SUCCESS] 成功读取 {len(df)} 道题目")
            return df

        except Exception as e:
            print(f"  [ERROR] 读取失败: {e}")
            return pd.DataFrame()

    def clean_data(self, df: pd.DataFrame) -> pd.DataFrame:
        """清理数据"""
        # 获取题型列名
        question_type_col = self.config["column_mapping"]["question_type"]

        # 检查该列是否存在
        if question_type_col not in df.columns:
            print(f"  [ERROR] 未找到题型列: '{question_type_col}'")
            print(f"  可用的列: {list(df.columns)}")
            return pd.DataFrame()

        # 移除空行
        df = df.dropna(subset=[question_type_col])

        # 移除重复的标题行
        question_type_col = self.config["column_mapping"]["question_type"]
        df = df[df[question_type_col] != question_type_col]

        # 标准化列名
        column_mapping = self.config["column_mapping"]

        # 创建新的DataFrame
        new_df = pd.DataFrame()

        # 必需的列
        required_columns = [
            "来源文件",
            column_mapping["question_type"],
            column_mapping["question_text"],
            column_mapping["correct_answer"]
        ]

        for col in required_columns:
            if col in df.columns:
                new_df[col] = df[col]

        # 可选的列
        optional_columns = [
            column_mapping["analysis"],
            column_mapping["score"],
            column_mapping["difficulty"]
        ]

        for col in optional_columns:
            if col in df.columns:
                new_df[col] = df[col]

        # 选项列
        for opt in column_mapping["options"]:
            if opt in df.columns:
                new_df[opt] = df[opt]

        return new_df

    def merge_files(self, input_dir: str = ".", file_pattern: str = None) -> pd.DataFrame:
        """合并所有Excel文件"""
        if file_pattern is None:
            # 尝试多个模式
            all_files = []
            for pattern in self.config["file_patterns"]:
                files = glob.glob(os.path.join(input_dir, pattern))
                all_files.extend(files)
            files = list(set(all_files))  # 去重
        else:
            files = glob.glob(os.path.join(input_dir, file_pattern))

        if not files:
            print("未找到任何Excel文件")
            return pd.DataFrame()

        print(f"找到 {len(files)} 个文件")

        all_data = []
        for file in sorted(files):
            data = self.read_excel_file(file)
            if not data.empty:
                all_data.append(data)

        if all_data:
            self.merged_data = pd.concat(all_data, ignore_index=True)
            print(f"\n成功合并 {len(self.merged_data)} 道题目")
            return self.merged_data
        else:
            print("没有成功读取任何文件")
            return pd.DataFrame()

    def save_excel(self, output_path: str = None):
        """保存为Excel文件"""
        if self.merged_data is None or self.merged_data.empty:
            print("没有数据可保存")
            return

        if output_path is None:
            output_path = self.config["output_settings"]["excel_filename"]

        # 创建输出目录
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        self.merged_data.to_excel(output_path, index=False, engine='openpyxl')
        print(f"[SUCCESS] Excel文件已保存: {output_path}")

    def save_word(self, output_path: str = None):
        """保存为Word文档"""
        if not DOCX_AVAILABLE:
            print("无法生成Word文档：请安装 python-docx")
            return

        if self.merged_data is None or self.merged_data.empty:
            print("没有数据可保存")
            return

        if output_path is None:
            output_path = self.config["output_settings"]["word_filename"]

        # 创建输出目录
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        # 创建Word文档
        doc = Document()

        # 标题
        title = doc.add_heading('题库汇总文档', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 统计信息
        doc.add_paragraph(f'总计 {len(self.merged_data)} 道题目\n')

        # 按来源分组
        current_source = None
        question_number = 0

        column_map = self.config["column_mapping"]

        for idx, row in self.merged_data.iterrows():
            question_number += 1

            # 新来源的标题
            if row["来源文件"] != current_source:
                current_source = row["来源文件"]
                doc.add_page_break()
                doc.add_heading(f'{current_source}', level=1)

            # 题目
            p = doc.add_paragraph()
            p.add_run(f'{question_number}. ').bold = True
            p.add_run(f'[{row[column_map["question_type"]]}] ')
            p.add_run(str(row[column_map["question_text"]]))

            # 选项
            for i, opt in enumerate(column_map["options"]):
                if opt in row and pd.notna(row[opt]):
                    p_opt = doc.add_paragraph()
                    p_opt.add_run(f'{chr(65+i)}. ').bold = True
                    p_opt.add_run(str(row[opt]))

            # 答案
            if pd.notna(row[column_map["correct_answer"]]):
                p_answer = doc.add_paragraph()
                p_answer.add_run('正确答案：').bold = True
                p_answer.add_run(str(row[column_map["correct_answer"]]))

            # 解析
            if (self.config["output_settings"]["include_analysis"] and
                column_map["analysis"] in row and
                pd.notna(row[column_map["analysis"]])):
                p_analysis = doc.add_paragraph()
                p_analysis.add_run('解析：').bold = True
                p_analysis.add_run(str(row[column_map["analysis"]]))

            doc.add_paragraph()  # 空行

        doc.save(output_path)
        print(f"[SUCCESS] Word文档已保存: {output_path}")

    def generate_report(self) -> Dict:
        """生成统计报告"""
        if self.merged_data is None:
            return {}

        report = {
            "总题目数": len(self.merged_data),
            "按来源统计": {},
            "按题型统计": {},
            "答案缺失统计": {}
        }

        # 按来源统计
        source_counts = self.merged_data["来源文件"].value_counts()
        report["按来源统计"] = source_counts.to_dict()

        # 按题型统计
        column_map = self.config["column_mapping"]
        if column_map["question_type"] in self.merged_data.columns:
            type_counts = self.merged_data[column_map["question_type"]].value_counts()
            report["按题型统计"] = type_counts.to_dict()

        # 答案缺失统计
        if column_map["correct_answer"] in self.merged_data.columns:
            missing_count = self.merged_data[column_map["correct_answer"]].isnull().sum()
            report["答案缺失统计"] = {
                "缺失数量": missing_count,
                "缺失比例": f"{missing_count/len(self.merged_data)*100:.1f}%"
            }

        return report


def main():
    parser = argparse.ArgumentParser(description="题库合并工具")
    parser.add_argument("--config", default="config/config.json", help="配置文件路径")
    parser.add_argument("--input", default=".", help="输入目录")
    parser.add_argument("--pattern", help="文件匹配模式")
    parser.add_argument("--output-excel", help="Excel输出文件路径")
    parser.add_argument("--output-word", help="Word输出文件路径")
    parser.add_argument("--word-only", action="store_true", help="只生成Word文档")
    parser.add_argument("--excel-only", action="store_true", help="只生成Excel文件")

    args = parser.parse_args()

    # 创建合并器
    merger = QuestionBankMerger(args.config)

    # 合并文件
    data = merger.merge_files(args.input, args.pattern)

    if data.empty:
        print("没有数据可处理")
        return

    # 生成报告
    report = merger.generate_report()
    print("\n=== 统计报告 ===")
    for key, value in report.items():
        print(f"{key}:")
        if isinstance(value, dict):
            for k, v in value.items():
                print(f"  {k}: {v}")
        else:
            print(f"  {value}")

    # 保存文件
    if not args.word_only:
        merger.save_excel(args.output_excel)

    if not args.excel_only:
        merger.save_word(args.output_word)


if __name__ == "__main__":
    main()